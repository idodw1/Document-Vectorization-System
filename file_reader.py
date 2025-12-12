"""
File Reader Module
Handles reading text from PDF and DOCX files
"""

from pathlib import Path
from typing import Tuple
from PyPDF2 import PdfReader
from docx import Document


class FileReader:
    """Class to handle reading different file types"""

    SUPPORTED_EXTENSIONS = {'.pdf', '.docx'}

    @staticmethod
    def get_file_type(filepath: str) -> str:
        """
        Get the file extension

        Args:
            filepath: Path to the file

        Returns:
            File extension (e.g., '.pdf', '.docx')

        Raises:
            ValueError: If file type is not supported
        """
        path = Path(filepath)
        if not path.exists():
            raise FileNotFoundError(f"File not found: {filepath}")

        extension = path.suffix.lower()
        if extension not in FileReader.SUPPORTED_EXTENSIONS:
            raise ValueError(
                f"Unsupported file type: {extension}. "
                f"Supported types: {FileReader.SUPPORTED_EXTENSIONS}"
            )

        return extension

    @staticmethod
    def read_pdf(filepath: str) -> str:
        """
        Extract text from PDF file

        Args:
            filepath: Path to PDF file

        Returns:
            Extracted text as string
        """
        try:
            reader = PdfReader(filepath)
            text_parts = []

            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)

            full_text = "\n".join(text_parts)
            print(f"✓ Extracted {len(full_text)} characters from {len(reader.pages)} pages")
            return full_text

        except Exception as e:
            raise Exception(f"Error reading PDF: {str(e)}")

    @staticmethod
    def read_docx(filepath: str) -> str:
        """
        Extract text from DOCX file

        Args:
            filepath: Path to DOCX file

        Returns:
            Extracted text as string
        """
        try:
            doc = Document(filepath)
            text_parts = []

            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            full_text = "\n".join(text_parts)
            print(f"✓ Extracted {len(full_text)} characters from {len(doc.paragraphs)} paragraphs")
            return full_text

        except Exception as e:
            raise Exception(f"Error reading DOCX: {str(e)}")

    @classmethod
    def read_file(cls, filepath: str) -> Tuple[str, str]:
        """
        Read file and return text content with filename

        Args:
            filepath: Path to the file

        Returns:
            Tuple of (text_content, filename)
        """
        file_type = cls.get_file_type(filepath)
        filename = Path(filepath).name

        print(f"Reading {file_type} file: {filename}")

        if file_type == '.pdf':
            text = cls.read_pdf(filepath)
        elif file_type == '.docx':
            text = cls.read_docx(filepath)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")

        return text, filename


if __name__ == "__main__":
    # Test the file reader
    import sys

    if len(sys.argv) < 2:
        print("Usage: python file_reader.py <filepath>")
        sys.exit(1)

    filepath = sys.argv[1]
    try:
        text, filename = FileReader.read_file(filepath)
        print(f"\n✓ Successfully read {filename}")
        print(f"Text preview (first 200 chars):\n{text[:200]}...")
    except Exception as e:
        print(f"✗ Error: {e}")
        sys.exit(1)